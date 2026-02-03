"""
Testes unitários para a Etapa 3 - Funcionalidades de Automação
"""

import pytest
from pathlib import Path
import tempfile
import shutil

from src.models.documento_rpcm import DocumentoRPCM
from src.converters.html_to_docx import HTMLtoDOCXConverter
from src.converters.word_html_cleaner import WordHTMLCleaner


class TestDocumentoRPCM:
    """Testes da classe DocumentoRPCM"""
    
    def test_criacao_valida(self):
        """Testa criação com dados válidos"""
        doc = DocumentoRPCM(
            grupo="INFRAESTRUTURA",
            subgrupo="ÁGUA",
            numero_preco="123456",
            descricao="Tubulação PVC 50mm",
            unidade="m",
            regulamentacao_html="<p>Regulamentação</p>"
        )
        assert doc.grupo == "INFRAESTRUTURA"
        assert doc.numero_preco == "123456"
    
    def test_subgrupo_vazio_valido(self):
        """Testa que subgrupo pode estar vazio"""
        doc = DocumentoRPCM(
            grupo="GRUPO",
            subgrupo="",  # VAZIO - deve ser aceito
            numero_preco="123456",
            descricao="DESC",
            unidade="m",
            regulamentacao_html="<p>R</p>"
        )
        assert doc.subgrupo == ""
    
    def test_validacao_grupo_vazio(self):
        """Testa rejeição de grupo vazio"""
        with pytest.raises(ValueError, match="Campos obrigatórios vazios"):
            DocumentoRPCM(
                grupo="",  # VAZIO - deve falhar
                subgrupo="SUB",
                numero_preco="123456",
                descricao="DESC",
                unidade="m",
                regulamentacao_html="<p>R</p>"
            )
    
    def test_validacao_numero_preco_vazio(self):
        """Testa rejeição de número preço vazio"""
        with pytest.raises(ValueError, match="Campos obrigatórios vazios"):
            DocumentoRPCM(
                grupo="GRUPO",
                subgrupo="",
                numero_preco="",  # VAZIO - deve falhar
                descricao="DESC",
                unidade="m",
                regulamentacao_html="<p>R</p>"
            )
    
    def test_validacao_descricao_vazia(self):
        """Testa rejeição de descrição vazia"""
        with pytest.raises(ValueError, match="Campos obrigatórios vazios"):
            DocumentoRPCM(
                grupo="GRUPO",
                subgrupo="",
                numero_preco="123456",
                descricao="",  # VAZIO - deve falhar
                unidade="m",
                regulamentacao_html="<p>R</p>"
            )
    
    def test_validacao_unidade_vazia(self):
        """Testa rejeição de unidade vazia"""
        with pytest.raises(ValueError, match="Campos obrigatórios vazios"):
            DocumentoRPCM(
                grupo="GRUPO",
                subgrupo="",
                numero_preco="123456",
                descricao="DESC",
                unidade="",  # VAZIO - deve falhar
                regulamentacao_html="<p>R</p>"
            )
    
    def test_validacao_regulamentacao_vazia(self):
        """Testa rejeição de regulamentação vazia"""
        with pytest.raises(ValueError, match="Campos obrigatórios vazios"):
            DocumentoRPCM(
                grupo="GRUPO",
                subgrupo="",
                numero_preco="123456",
                descricao="DESC",
                unidade="m",
                regulamentacao_html=""  # VAZIO - deve falhar
            )
    
    def test_limpeza_espacos(self):
        """Testa remoção de espaços extras"""
        doc = DocumentoRPCM(
            grupo="  GRUPO  ",
            subgrupo=" SUB ",
            numero_preco="  123456  ",
            descricao="  DESC  ",
            unidade=" m ",
            regulamentacao_html="<p>R</p>"
        )
        assert doc.grupo == "GRUPO"
        assert doc.subgrupo == "SUB"
        assert doc.numero_preco == "123456"
        assert doc.descricao == "DESC"
        assert doc.unidade == "m"
    
    def test_nome_arquivo_simples(self):
        """Testa geração de nome de arquivo simples"""
        doc = DocumentoRPCM(
            grupo="G", subgrupo="S", numero_preco="123456",
            descricao="Teste Simples", unidade="m",
            regulamentacao_html="<p>R</p>"
        )
        assert doc.get_nome_arquivo() == "123456_Teste Simples.docx"
    
    def test_nome_arquivo_caracteres_invalidos(self):
        """Testa remoção de caracteres inválidos"""
        doc = DocumentoRPCM(
            grupo="G", subgrupo="", numero_preco="789012",
            descricao="Teste/Com\\Caracteres:Inválidos*?<>|",
            unidade="m", regulamentacao_html="<p>R</p>"
        )
        nome = doc.get_nome_arquivo()
        caracteres_invalidos = ['/', '\\', ':', '*', '?', '<', '>', '|']
        for char in caracteres_invalidos:
            assert char not in nome
    
    def test_nome_arquivo_longo(self):
        """Testa truncamento de nome muito longo"""
        descricao_longa = "A" * 150
        doc = DocumentoRPCM(
            grupo="G", subgrupo="", numero_preco="999999",
            descricao=descricao_longa, unidade="m",
            regulamentacao_html="<p>R</p>"
        )
        nome = doc.get_nome_arquivo()
        # 100 chars descrição + 6 numero + 1 underscore + 5 .docx = 112
        assert len(nome) <= 112
    
    def test_to_dict(self):
        """Testa conversão para dicionário"""
        doc = DocumentoRPCM(
            grupo="G", subgrupo="S", numero_preco="01",
            descricao="D", unidade="m",
            regulamentacao_html="<p>R</p>"
        )
        d = doc.to_dict()
        assert d['GRUPO'] == "G"
        assert d['SUBGRUPO'] == "S"
        assert d['N_PRECO'] == "01"
        assert d['DESCRICAO'] == "D"
        assert d['UNIDADE'] == "m"


class TestHTMLtoDOCXConverter:
    """Testes do conversor HTML → DOCX"""
    
    def test_paragrafo_simples(self):
        """Testa conversão de parágrafo simples"""
        html = "<p>Texto simples</p>"
        converter = HTMLtoDOCXConverter()
        doc = converter.convert(html)
        
        assert len(doc.paragraphs) > 0
        assert "Texto simples" in doc.paragraphs[0].text
    
    def test_texto_negrito(self):
        """Testa preservação de negrito"""
        html = "<p>Texto <strong>negrito</strong> normal</p>"
        converter = HTMLtoDOCXConverter()
        doc = converter.convert(html)
        
        # Verificar que existe run com negrito
        runs = doc.paragraphs[0].runs
        tem_negrito = any(run.bold for run in runs if run.text.strip())
        assert tem_negrito
    
    def test_texto_italico(self):
        """Testa preservação de itálico"""
        html = "<p>Texto <em>itálico</em> normal</p>"
        converter = HTMLtoDOCXConverter()
        doc = converter.convert(html)
        
        runs = doc.paragraphs[0].runs
        tem_italico = any(run.italic for run in runs if run.text.strip())
        assert tem_italico
    
    def test_texto_sublinhado(self):
        """Testa preservação de sublinhado"""
        html = "<p>Texto <u>sublinhado</u> normal</p>"
        converter = HTMLtoDOCXConverter()
        doc = converter.convert(html)
        
        runs = doc.paragraphs[0].runs
        tem_sublinhado = any(run.underline for run in runs if run.text.strip())
        assert tem_sublinhado
    
    def test_lista_nao_ordenada(self):
        """Testa conversão de lista com marcadores"""
        html = """
        <ul>
            <li>Item 1</li>
            <li>Item 2</li>
            <li>Item 3</li>
        </ul>
        """
        converter = HTMLtoDOCXConverter()
        doc = converter.convert(html)
        
        # Verificar que foram criados parágrafos
        assert len(doc.paragraphs) >= 3
    
    def test_lista_ordenada(self):
        """Testa conversão de lista numerada"""
        html = """
        <ol>
            <li>Primeiro</li>
            <li>Segundo</li>
            <li>Terceiro</li>
        </ol>
        """
        converter = HTMLtoDOCXConverter()
        doc = converter.convert(html)
        
        assert len(doc.paragraphs) >= 3
    
    def test_tabela_simples(self):
        """Testa conversão de tabela"""
        html = """
        <table>
            <tr>
                <th>Cabeçalho 1</th>
                <th>Cabeçalho 2</th>
            </tr>
            <tr>
                <td>Célula 1</td>
                <td>Célula 2</td>
            </tr>
        </table>
        """
        converter = HTMLtoDOCXConverter()
        doc = converter.convert(html)
        
        assert len(doc.tables) >= 1
        assert len(doc.tables[0].rows) == 2
        assert len(doc.tables[0].columns) == 2
    
    def test_alinhamento_justificado(self):
        """Testa alinhamento justificado"""
        html = '<p style="text-align: justify">Texto justificado</p>'
        converter = HTMLtoDOCXConverter()
        doc = converter.convert(html)
        
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        assert doc.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    
    def test_espacamento_15(self):
        """Testa que espaçamento 1,5 é aplicado"""
        html = "<p>Texto com espaçamento</p>"
        converter = HTMLtoDOCXConverter()
        doc = converter.convert(html)
        
        # Verificar espaçamento
        para = doc.paragraphs[0]
        assert para.paragraph_format.line_spacing == 1.5


class TestWordHTMLCleaner:
    """Testes do limpador de HTML do Word"""
    
    def test_remover_tags_word(self):
        """Testa remoção de tags específicas do Word"""
        html = '<p>Texto<o:p></o:p> normal</p>'
        limpo = WordHTMLCleaner.clean(html)
        assert '<o:p>' not in limpo
        assert 'Texto' in limpo
    
    def test_remover_classes_mso(self):
        """Testa remoção de classes Mso"""
        html = '<p class="MsoNormal">Texto</p>'
        limpo = WordHTMLCleaner.clean(html)
        assert 'MsoNormal' not in limpo
    
    def test_manter_formatacao_essencial(self):
        """Testa que mantém formatações importantes"""
        html = '<p style="font-size: 12pt; font-weight: bold">Texto</p>'
        limpo = WordHTMLCleaner.clean(html)
        assert 'font-size' in limpo
        assert 'font-weight' in limpo


if __name__ == '__main__':
    pytest.main([__file__, '-v'])
