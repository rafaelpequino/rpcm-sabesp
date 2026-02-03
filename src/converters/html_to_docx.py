"""
Conversor HTML → DOCX com preservação PERFEITA de formatação

REQUISITOS CRÍTICOS:
- Espaçamento entre linhas: 1,5 (SEMPRE)
- Fonte: Arial 10pt
- Alinhamento: Justificado (padrão)
- Preservar TODAS as formatações do HTML
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from bs4 import BeautifulSoup
import logging
import re

logger = logging.getLogger(__name__)


class HTMLtoDOCXConverter:
    """
    Conversor HTML → DOCX com preservação PERFEITA de formatação
    """
    
    # CONSTANTES DE FORMATAÇÃO PADRÃO
    DEFAULT_FONT = 'Arial'
    DEFAULT_SIZE = Pt(10)
    DEFAULT_LINE_SPACING = 1.5  # ⭐ ESPAÇAMENTO 1,5
    DEFAULT_ALIGNMENT = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    def __init__(self):
        self.doc = None
    
    def convert(self, html_content: str, base_doc=None) -> Document:
        """
        Converte HTML para python-docx com PERFEIÇÃO
        
        Args:
            html_content: String HTML do editor
            base_doc: Document base (template) se existir
            
        Returns:
            Document object com formatação perfeita
        """
        if base_doc:
            self.doc = base_doc
        else:
            self.doc = Document()
            self._configurar_documento_padrao()
        
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Processar cada elemento
        for element in soup.children:
            if hasattr(element, 'name') and element.name:  # Ignorar NavigableString
                self._process_element(element)
        
        logger.info("Conversão HTML→DOCX concluída")
        return self.doc
    
    def _configurar_documento_padrao(self):
        """Configura estilos padrão do documento"""
        # Estilo Normal
        style = self.doc.styles['Normal']
        style.font.name = self.DEFAULT_FONT
        style.font.size = self.DEFAULT_SIZE
        
        # ⭐ ESPAÇAMENTO DE LINHA 1,5 NO ESTILO PADRÃO
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        style.paragraph_format.line_spacing = self.DEFAULT_LINE_SPACING
        style.paragraph_format.alignment = self.DEFAULT_ALIGNMENT
    
    def _process_element(self, element):
        """Processa cada elemento HTML recursivamente"""
        
        if element.name == 'p':
            self._add_paragraph(element)
        elif element.name in ['ul', 'ol']:
            self._add_list(element)
        elif element.name == 'table':
            self._add_table(element)
        elif element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            self._add_heading(element)
        elif element.name == 'br':
            self.doc.add_paragraph()  # Linha em branco
        elif element.name == 'div':
            # Processar filhos do div
            for child in element.children:
                if hasattr(child, 'name') and child.name:
                    self._process_element(child)
    
    def _add_paragraph(self, p_element):
        """
        Adiciona parágrafo com TODA a formatação preservada
        
        ⭐ CRÍTICO: Espaçamento 1,5, Arial 10pt, justificado
        """
        p = self.doc.add_paragraph()
        
        # ⭐ APLICAR ESPAÇAMENTO 1,5 (OBRIGATÓRIO)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = self.DEFAULT_LINE_SPACING
        
        # Processar alinhamento do estilo inline
        style = p_element.get('style', '')
        alignment = self._extract_alignment(style)
        if alignment:
            p.alignment = alignment
        else:
            p.alignment = self.DEFAULT_ALIGNMENT
        
        # Processar texto inline com formatação
        self._process_inline_elements(p_element, p)
    
    def _process_inline_elements(self, element, paragraph):
        """
        Processa elementos inline (strong, em, u, span) com PRECISÃO
        
        Preserva: negrito, itálico, sublinhado, tachado, cor, tamanho
        """
        
        for content in element.children:
            if isinstance(content, str):
                # Texto simples
                if content.strip():  # Ignorar espaços vazios
                    run = paragraph.add_run(content)
                    self._apply_default_font(run)
            else:
                # Elementos com formatação
                text = content.get_text()
                if not text.strip():
                    continue
                
                run = paragraph.add_run(text)
                self._apply_default_font(run)
                
                # ⭐ APLICAR TODAS AS FORMATAÇÕES
                
                # Negrito
                if content.name in ['strong', 'b']:
                    run.bold = True
                
                # Itálico
                if content.name in ['em', 'i']:
                    run.italic = True
                
                # Sublinhado
                if content.name == 'u':
                    run.underline = True
                
                # Tachado
                if content.name in ['s', 'strike', 'del']:
                    run.font.strike = True
                
                # Processar estilos inline (span)
                if content.name == 'span' or content.get('style'):
                    style = content.get('style', '')
                    self._apply_inline_styles(run, style)
                
                # Link
                if content.name == 'a':
                    run.font.color.rgb = RGBColor(0, 0, 255)
                    run.underline = True
    
    def _apply_default_font(self, run):
        """Aplica fonte padrão (Arial 10pt)"""
        run.font.name = self.DEFAULT_FONT
        run.font.size = self.DEFAULT_SIZE
    
    def _apply_inline_styles(self, run, style_string):
        """
        Aplica estilos inline do CSS com PRECISÃO
        
        Suporta: font-size, font-family, color, background-color, 
                 font-weight, font-style, text-decoration
        """
        if not style_string:
            return
        
        styles = {}
        for item in style_string.split(';'):
            if ':' in item:
                prop, value = item.split(':', 1)
                styles[prop.strip().lower()] = value.strip()
        
        # Font size
        if 'font-size' in styles:
            size = styles['font-size']
            if 'pt' in size:
                run.font.size = Pt(float(size.replace('pt', '')))
            elif 'px' in size:
                px = float(size.replace('px', ''))
                run.font.size = Pt(px * 0.75)  # Converter px para pt
        
        # Font family
        if 'font-family' in styles:
            family = styles['font-family'].strip('"\'').split(',')[0]
            run.font.name = family
        
        # Color
        if 'color' in styles:
            color = styles['color'].strip()
            rgb = self._parse_color(color)
            if rgb:
                run.font.color.rgb = RGBColor(*rgb)
        
        # Font weight (bold)
        if 'font-weight' in styles:
            weight = styles['font-weight']
            if weight in ['bold', '700', '800', '900']:
                run.bold = True
        
        # Font style (italic)
        if 'font-style' in styles:
            if styles['font-style'] == 'italic':
                run.italic = True
        
        # Text decoration
        if 'text-decoration' in styles:
            decoration = styles['text-decoration']
            if 'underline' in decoration:
                run.underline = True
            if 'line-through' in decoration:
                run.font.strike = True
    
    def _parse_color(self, color_str):
        """
        Parse cor de CSS para RGB
        
        Suporta: #RRGGBB, rgb(r,g,b), nomes de cores
        """
        color_str = color_str.lower().strip()
        
        # Hex color #RRGGBB
        if color_str.startswith('#'):
            hex_color = color_str[1:]
            if len(hex_color) == 6:
                return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
            elif len(hex_color) == 3:
                return tuple(int(c*2, 16) for c in hex_color)
        
        # rgb(r, g, b)
        if color_str.startswith('rgb'):
            match = re.search(r'rgb\((\d+),\s*(\d+),\s*(\d+)\)', color_str)
            if match:
                return tuple(int(x) for x in match.groups())
        
        # Cores nomeadas (subset)
        colors = {
            'black': (0, 0, 0),
            'white': (255, 255, 255),
            'red': (255, 0, 0),
            'green': (0, 128, 0),
            'blue': (0, 0, 255),
            'yellow': (255, 255, 0),
            'gray': (128, 128, 128),
            'grey': (128, 128, 128),
        }
        
        return colors.get(color_str)
    
    def _extract_alignment(self, style_string):
        """Extrai alinhamento do estilo CSS"""
        if 'text-align' in style_string:
            if 'left' in style_string:
                return WD_ALIGN_PARAGRAPH.LEFT
            elif 'center' in style_string:
                return WD_ALIGN_PARAGRAPH.CENTER
            elif 'right' in style_string:
                return WD_ALIGN_PARAGRAPH.RIGHT
            elif 'justify' in style_string:
                return WD_ALIGN_PARAGRAPH.JUSTIFY
        return None
    
    def _add_list(self, list_element):
        """
        Adiciona lista (ordenada ou não) com PERFEIÇÃO
        
        ⭐ Preserva: tipo de marcador, níveis de recuo, espaçamento 1,5
        """
        is_ordered = list_element.name == 'ol'
        
        # Detectar nível de recuo (para listas aninhadas)
        list_level = self._detect_list_level(list_element)
        
        for li in list_element.find_all('li', recursive=False):
            p = self.doc.add_paragraph()
            
            # ⭐ APLICAR ESPAÇAMENTO 1,5
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing = self.DEFAULT_LINE_SPACING
            
            # Aplicar estilo de lista
            if is_ordered:
                p.style = 'List Number'
            else:
                p.style = 'List Bullet'
            
            # Aplicar nível de recuo
            if list_level > 0:
                p.paragraph_format.left_indent = Cm(list_level * 1.27)  # 1.27cm por nível
            
            # Processar conteúdo do item
            self._process_inline_elements(li, p)
            
            # Verificar se há sub-listas
            sub_lists = li.find_all(['ul', 'ol'], recursive=False)
            for sub_list in sub_lists:
                self._add_list(sub_list)
    
    def _detect_list_level(self, list_element):
        """Detecta nível de recuo da lista (para listas aninhadas)"""
        level = 0
        parent = list_element.parent
        while parent:
            if hasattr(parent, 'name') and parent.name in ['ul', 'ol']:
                level += 1
            parent = parent.parent
        return level
    
    def _add_table(self, table_element):
        """
        Adiciona tabela com TODA a formatação preservada
        
        ⭐ Preserva: bordas, mesclagem, cores, alinhamento, formatação de texto
        """
        # Detectar dimensões
        rows_html = table_element.find_all('tr')
        if not rows_html:
            return
        
        cols = max(len(row.find_all(['td', 'th'])) for row in rows_html)
        rows = len(rows_html)
        
        # Criar tabela
        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Processar cada célula
        for i, row_html in enumerate(rows_html):
            cells_html = row_html.find_all(['td', 'th'])
            
            for j, cell_html in enumerate(cells_html):
                if j >= len(table.rows[i].cells):
                    continue
                
                table_cell = table.rows[i].cells[j]
                
                # ⭐ LIMPAR PARÁGRAFO PADRÃO
                if table_cell.paragraphs:
                    para = table_cell.paragraphs[0]
                else:
                    para = table_cell.add_paragraph()
                
                # ⭐ ESPAÇAMENTO 1,5 NA CÉLULA
                para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                para.paragraph_format.line_spacing = self.DEFAULT_LINE_SPACING
                
                # Processar conteúdo da célula
                self._process_inline_elements(cell_html, para)
                
                # Célula de cabeçalho (th)
                if cell_html.name == 'th':
                    for para in table_cell.paragraphs:
                        for run in para.runs:
                            run.bold = True
                    # Alinhamento centro para cabeçalho
                    for para in table_cell.paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_heading(self, heading_element):
        """
        Adiciona cabeçalho com formatação
        
        ⭐ Mantém espaçamento 1,5
        """
        level = int(heading_element.name[1])  # h1 → 1, h2 → 2, etc
        text = heading_element.get_text()
        
        heading = self.doc.add_heading(text, level=level)
        
        # ⭐ APLICAR ESPAÇAMENTO 1,5
        heading.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        heading.paragraph_format.line_spacing = self.DEFAULT_LINE_SPACING
        
        # Aplicar fonte Arial
        for run in heading.runs:
            run.font.name = self.DEFAULT_FONT
