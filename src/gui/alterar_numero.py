"""
Módulo para Alterar Números em Arquivos Word
Substitui números em documentos DOCX
"""

import customtkinter as ctk
from tkinter import filedialog, messagebox
import threading
from pathlib import Path
import traceback
import os
from datetime import datetime

from src.gui.styles import COLORS, FONTS, SPACING

# Importar biblioteca de manipulação de DOCX
try:
    from docx import Document
    from docx.shared import RGBColor, Pt
    DOCX_DISPONIVEL = True
except ImportError:
    DOCX_DISPONIVEL = False


class AlterarNumeroFrame(ctk.CTkScrollableFrame):
    """Frame para alterar números em arquivos Word"""
    
    def __init__(self, parent):
        super().__init__(
            parent,
            fg_color="transparent",
            scrollbar_button_color=COLORS['primary'],
            scrollbar_button_hover_color=COLORS['hover']
        )
        
        # Configurar scroll suave
        self._parent_canvas.configure(yscrollincrement=20)
        self._configurar_scroll_suave()
        
        # Variáveis
        self.arquivo_word = None
        self.numero_busca = None
        self.numero_substituir = None
        self.pasta_destino = None
        self.processamento_ativo = False
        self.ocorrencias_encontradas = 0
        
        self._criar_interface()
    
    def _configurar_scroll_suave(self):
        """Configura scroll suave com mouse wheel"""
        def _on_mousewheel(event):
            self._parent_canvas.yview_scroll(int(-1 * (event.delta / 60)), "units")
        
        def _unbind_mousewheel(event):
            self._parent_canvas.unbind_all("<MouseWheel>")
        
        def _bind_mousewheel(event):
            self._parent_canvas.bind_all("<MouseWheel>", _on_mousewheel)
        
        self._parent_canvas.bind("<Leave>", _unbind_mousewheel)
        self._parent_canvas.bind("<Enter>", _bind_mousewheel)
    
    def _criar_interface(self):
        """Cria a interface da aba"""
        
        # Título
        titulo = ctk.CTkLabel(
            self,
            text="🔢 Alterar Nº Preço em massa",
            font=FONTS['title'],
            text_color=COLORS['primary']
        )
        titulo.pack(pady=(0, SPACING['padding']))
        
        # Descrição
        descricao = ctk.CTkLabel(
            self,
            text="Selecione um arquivo Word, informe o número a buscar e o número de substituição",
            font=FONTS['small'],
            text_color=COLORS['text_secondary']
        )
        descricao.pack(pady=(0, SPACING['padding']))
        
        # ===== SEÇÃO 1: SELEÇÃO DE ARQUIVO =====
        self._criar_secao_selecao_arquivo()
        
        # ===== SEÇÃO 2: NÚMEROS =====
        self._criar_secao_numeros()
        
        # ===== SEÇÃO 3: PASTA DE DESTINO =====
        self._criar_secao_pasta_destino()
        
        # ===== SEÇÃO 4: AÇÕES =====
        self._criar_secao_acoes()
    
    def _criar_secao_selecao_arquivo(self):
        """Cria seção de seleção de arquivo"""
        frame = ctk.CTkFrame(self, fg_color=COLORS['background'])
        frame.pack(fill="x", pady=SPACING['margin'])
        
        titulo_secao = ctk.CTkLabel(
            frame,
            text="Selecionar Arquivo Word",
            font=FONTS['subtitle'],
            text_color=COLORS['primary']
        )
        titulo_secao.pack(anchor="w", padx=SPACING['padding'], pady=(SPACING['padding'], SPACING['margin']))
        
        # Frame com botão e label
        container = ctk.CTkFrame(frame, fg_color="transparent")
        container.pack(fill="x", padx=SPACING['padding'], pady=(0, SPACING['padding']))
        
        self.btn_selecionar = ctk.CTkButton(
            container,
            text="📂 Selecionar Arquivo DOCX",
            command=self._selecionar_arquivo,
            width=250,
            height=40,
            font=FONTS['button'],
            fg_color=COLORS['primary'],
            hover_color=COLORS['hover']
        )
        self.btn_selecionar.pack(side="left", padx=(0, SPACING['margin']))
        
        # Label com nome do arquivo
        self.label_arquivo = ctk.CTkLabel(
            container,
            text="Nenhum arquivo selecionado",
            font=FONTS['input'],
            text_color=COLORS['text_secondary']
        )
        self.label_arquivo.pack(side="left", fill="x", expand=True)
    
    def _criar_secao_numeros(self):
        """Cria seção de entrada dos números"""
        frame = ctk.CTkFrame(self, fg_color=COLORS['background'])
        frame.pack(fill="x", pady=SPACING['margin'])
        
        titulo_secao = ctk.CTkLabel(
            frame,
            text="Definir Números",
            font=FONTS['subtitle'],
            text_color=COLORS['primary']
        )
        titulo_secao.pack(anchor="w", padx=SPACING['padding'], pady=(SPACING['padding'], SPACING['margin']))
        
        # Container para os inputs
        container_numeros = ctk.CTkFrame(frame, fg_color="transparent")
        container_numeros.pack(fill="x", padx=SPACING['padding'], pady=(0, SPACING['padding']))
        
        # Número a buscar
        label_busca = ctk.CTkLabel(
            container_numeros,
            text="Número a Buscar:",
            font=FONTS['label'],
            text_color=COLORS['text']
        )
        label_busca.grid(row=0, column=0, sticky="w", padx=(0, SPACING['margin']), pady=(0, SPACING['margin']))
        
        self.entrada_busca = ctk.CTkEntry(
            container_numeros,
            placeholder_text="Ex: 1234",
            font=FONTS['input'],
            width=200,
            height=SPACING['field_height']
        )
        self.entrada_busca.grid(row=0, column=1, sticky="w", padx=(0, SPACING['margin'] * 3))
        
        # Seta
        seta = ctk.CTkLabel(
            container_numeros,
            text="→",
            font=("Arial", 24),
            text_color=COLORS['primary']
        )
        seta.grid(row=0, column=2, padx=SPACING['margin'])
        
        # Número para substituir
        label_substituir = ctk.CTkLabel(
            container_numeros,
            text="Número Novo:",
            font=FONTS['label'],
            text_color=COLORS['text']
        )
        label_substituir.grid(row=0, column=3, sticky="w", padx=(SPACING['margin'], SPACING['margin']))
        
        self.entrada_substituir = ctk.CTkEntry(
            container_numeros,
            placeholder_text="Ex: 4321",
            font=FONTS['input'],
            width=200,
            height=SPACING['field_height']
        )
        self.entrada_substituir.grid(row=0, column=4, sticky="w")
        
        # Adicionar validação quando o usuário digita
        self.entrada_busca.bind("<KeyRelease>", lambda e: self._validar_entrada())
        self.entrada_substituir.bind("<KeyRelease>", lambda e: self._validar_entrada())
        
        # Configurar pesos das colunas para distribuição melhor
        container_numeros.grid_columnconfigure(1, minsize=200)
        container_numeros.grid_columnconfigure(4, minsize=200)
    
    def _criar_secao_pasta_destino(self):
        """Cria seção de seleção da pasta de destino"""
        frame = ctk.CTkFrame(self, fg_color=COLORS['background'])
        frame.pack(fill="x", pady=SPACING['margin'])
        
        titulo_secao = ctk.CTkLabel(
            frame,
            text="Pasta de Destino",
            font=FONTS['subtitle'],
            text_color=COLORS['primary']
        )
        titulo_secao.pack(anchor="w", padx=SPACING['padding'], pady=(SPACING['padding'], SPACING['margin']))
        
        # Frame com botão e label
        container = ctk.CTkFrame(frame, fg_color="transparent")
        container.pack(fill="x", padx=SPACING['padding'], pady=(0, SPACING['padding']))
        
        self.btn_selecionar_pasta = ctk.CTkButton(
            container,
            text="📁 Escolher Pasta",
            command=self._selecionar_pasta_destino,
            width=250,
            height=40,
            font=FONTS['button'],
            fg_color=COLORS['primary'],
            hover_color=COLORS['hover']
        )
        self.btn_selecionar_pasta.pack(side="left", padx=(0, SPACING['margin']))
        
        # Label com nome da pasta
        self.label_pasta = ctk.CTkLabel(
            container,
            text="Nenhuma pasta selecionada",
            font=FONTS['input'],
            text_color=COLORS['text_secondary']
        )
        self.label_pasta.pack(side="left", fill="x", expand=True)
    
    def _criar_secao_acoes(self):
        """Cria seção de botões de ação"""
        frame = ctk.CTkFrame(self, fg_color=COLORS['background'])
        frame.pack(fill="x", pady=SPACING['margin'])
        
        titulo_secao = ctk.CTkLabel(
            frame,
            text="Processar",
            font=FONTS['subtitle'],
            text_color=COLORS['primary']
        )
        titulo_secao.pack(anchor="w", padx=SPACING['padding'], pady=(SPACING['padding'], SPACING['margin']))
        
        container = ctk.CTkFrame(frame, fg_color="transparent")
        container.pack(fill="x", padx=SPACING['padding'], pady=(0, SPACING['padding']))
        
        # Botão processar
        self.btn_processar = ctk.CTkButton(
            container,
            text="🔄 Processar",
            command=self._processar_arquivo,
            width=250,
            height=45,
            font=FONTS['button'],
            fg_color=COLORS['success'],
            hover_color=('#22863A' if COLORS['success'] == '#28A745' else COLORS['success']),
            state="disabled"
        )
        self.btn_processar.pack(side="left", padx=(0, SPACING['margin']))
        
        # Informação
        self.label_info = ctk.CTkLabel(
            container,
            text="Selecione um arquivo, preencha os números e escolha a pasta de destino",
            font=FONTS['small'],
            text_color=COLORS['text_secondary']
        )
        self.label_info.pack(side="left", fill="x", expand=True)
    
    def _selecionar_pasta_destino(self):
        """Seleciona a pasta onde salvar o arquivo"""
        pasta = filedialog.askdirectory(
            title="Selecionar Pasta de Destino",
            parent=self
        )
        
        if pasta:
            # Validar se é a mesma pasta do arquivo original
            if self.arquivo_word:
                pasta_original = Path(self.arquivo_word).parent
                pasta_selecionada = Path(pasta)
                
                if pasta_original == pasta_selecionada:
                    messagebox.showwarning(
                        "Aviso",
                        "Não é possível salvar na mesma pasta do arquivo original.\n"
                        "Por favor, selecione uma pasta diferente."
                    )
                    return
            
            self.pasta_destino = pasta
            # Mostrar apenas o nome da pasta
            nome_pasta = Path(pasta).name
            self.label_pasta.configure(text=f"✓ {nome_pasta}")
            self._validar_entrada()
    
    def _selecionar_arquivo(self):
        """Seleciona um arquivo Word"""
        arquivo = filedialog.askopenfilename(
            title="Selecionar Arquivo Word",
            filetypes=[("Documentos Word", "*.docx"), ("Todos os arquivos", "*.*")],
            parent=self
        )
        
        if arquivo:
            self.arquivo_word = arquivo
            nome_arquivo = Path(arquivo).name
            self.label_arquivo.configure(text=f"✓ {nome_arquivo}")
            self._validar_entrada()
            self.label_info.configure(text=f"Arquivo: {nome_arquivo}")
    
    def _validar_entrada(self):
        """Valida se todos os campos estão preenchidos"""
        arquivo_ok = self.arquivo_word is not None
        numero_busca_ok = self.entrada_busca.get().strip() != ""
        numero_substituir_ok = self.entrada_substituir.get().strip() != ""
        pasta_ok = self.pasta_destino is not None
        pastas_diferentes = True
        
        # Verificar se a pasta de destino é diferente da pasta do arquivo original
        if arquivo_ok and pasta_ok:
            pasta_original = Path(self.arquivo_word).parent
            pasta_destino = Path(self.pasta_destino)
            if pasta_original == pasta_destino:
                pastas_diferentes = False
        
        if arquivo_ok and numero_busca_ok and numero_substituir_ok and pasta_ok and pastas_diferentes:
            self.btn_processar.configure(state="normal")
        else:
            self.btn_processar.configure(state="disabled")
    
    def _processar_arquivo(self):
        """Processa o arquivo em uma thread separada"""
        if not DOCX_DISPONIVEL:
            messagebox.showerror(
                "Erro",
                "Biblioteca python-docx não está instalada.\nPor favor, instale-a com: pip install python-docx"
            )
            return
        
        # Validar entrada
        numero_busca = self.entrada_busca.get().strip()
        numero_substituir = self.entrada_substituir.get().strip()
        
        if not numero_busca or not numero_substituir:
            messagebox.showwarning("Aviso", "Por favor, preencha ambos os números")
            return
        
        if numero_busca == numero_substituir:
            messagebox.showwarning("Aviso", "Os números não podem ser iguais")
            return
        
        # Iniciar processamento em thread
        self.processamento_ativo = True
        self.btn_processar.configure(state="disabled", text="⏳ Processando...")
        thread = threading.Thread(
            target=self._executar_substituicao,
            args=(self.arquivo_word, numero_busca, numero_substituir),
            daemon=True
        )
        thread.start()
    
    def _executar_substituicao(self, arquivo, numero_busca, numero_substituir):
        """Executa a substituição de números no documento"""
        try:
            # Carregar documento
            doc = Document(arquivo)
            
            count = 0
            
            # Procurar e substituir em parágrafos
            for paragrafo in doc.paragraphs:
                if numero_busca in paragrafo.text:
                    # Processar runs para manter formatação
                    for run in paragrafo.runs:
                        if numero_busca in run.text:
                            run.text = run.text.replace(numero_busca, numero_substituir)
                            count += run.text.count(numero_substituir)
            
            # Procurar e substituir em tabelas
            for tabela in doc.tables:
                for linha in tabela.rows:
                    for celula in linha.cells:
                        for paragrafo in celula.paragraphs:
                            if numero_busca in paragrafo.text:
                                for run in paragrafo.runs:
                                    if numero_busca in run.text:
                                        run.text = run.text.replace(numero_busca, numero_substituir)
                                        count += run.text.count(numero_substituir)
            
            # Salvar arquivo modificado na pasta de destino com o mesmo nome
            nome_arquivo_original = Path(self.arquivo_word).name
            arquivo_destino = Path(self.pasta_destino) / nome_arquivo_original
            
            doc.save(str(arquivo_destino))
            
            self.ocorrencias_encontradas = count
            
            # Atualizar interface na thread principal
            self.after(0, self._procesamento_concluido, count, str(arquivo_destino))
            
        except Exception as e:
            self.after(0, lambda: messagebox.showerror(
                "Erro na Substituição",
                f"Erro ao processar arquivo:\n{str(e)}\n\n{traceback.format_exc()}"
            ))
            self.after(0, self._procesamento_erro)
    
    def _procesamento_concluido(self, count, arquivo_destino):
        """Atualiza interface após sucesso"""
        self.processamento_ativo = False
        self.btn_processar.configure(state="normal", text="🔄 Processar")
        
        messagebox.showinfo(
            "Sucesso",
            f"Arquivo processado com sucesso!\n\n"
            f"Ocorrências alteradas: {count}\n\n"
            f"Arquivo salvo em:\n{arquivo_destino}"
        )
    
    def _procesamento_erro(self):
        """Atualiza interface após erro"""
        self.processamento_ativo = False
        self.btn_processar.configure(state="normal", text="🔄 Processar")
