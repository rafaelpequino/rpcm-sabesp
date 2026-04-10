"""
Módulo para Alterar Números em Múltiplos Arquivos Word (Processamento em Lote)
Usa uma planilha para fazer substituições em massa
"""

import customtkinter as ctk
from tkinter import filedialog, messagebox
import threading
from pathlib import Path
import traceback
import os
from datetime import datetime

from src.gui.styles import COLORS, FONTS, SPACING

# Importar bibliotecas necessárias
try:
    from docx import Document
    DOCX_DISPONIVEL = True
except ImportError:
    DOCX_DISPONIVEL = False

try:
    import pandas as pd
    PANDAS_DISPONIVEL = True
except ImportError:
    PANDAS_DISPONIVEL = False

try:
    from openpyxl import load_workbook
    OPENPYXL_DISPONIVEL = True
except ImportError:
    OPENPYXL_DISPONIVEL = False


class AlterarNumeroLoteFrame(ctk.CTkScrollableFrame):
    """Frame para alterar números em múltiplos arquivos Word"""
    
    def __init__(self, parent):
        super().__init__(
            parent,
            fg_color="transparent",
            scrollbar_button_color=COLORS['primary'],
            scrollbar_button_hover_color=COLORS['hover']
        )
        
        # Configurar scroll suave
        self._parent_canvas.configure(yscrollincrement=20)
        self._configurar_scroll_suave()
        
        # Variáveis
        self.arquivo_planilha = None
        self.pasta_arquivos = None
        self.processamento_ativo = False
        self.dados_planilha = []
        
        self._criar_interface()
    
    def _configurar_scroll_suave(self):
        """Configura scroll suave com mouse wheel"""
        def _on_mousewheel(event):
            self._parent_canvas.yview_scroll(int(-1 * (event.delta / 60)), "units")
        
        def _unbind_mousewheel(event):
            self._parent_canvas.unbind_all("<MouseWheel>")
        
        def _bind_mousewheel(event):
            self._parent_canvas.bind_all("<MouseWheel>", _on_mousewheel)
        
        self._parent_canvas.bind("<Leave>", _unbind_mousewheel)
        self._parent_canvas.bind("<Enter>", _bind_mousewheel)
    
    def _criar_interface(self):
        """Cria a interface da aba"""
        
        # Título
        titulo = ctk.CTkLabel(
            self,
            text="🔢 Alterar Nº Preço",
            font=FONTS['title'],
            text_color=COLORS['primary']
        )
        titulo.pack(pady=(0, SPACING['padding']))
        
        # Descrição
        descricao = ctk.CTkLabel(
            self,
            text="Processe múltiplos arquivos Word de uma vez usando uma planilha com os mapeamentos",
            font=FONTS['small'],
            text_color=COLORS['text_secondary']
        )
        descricao.pack(pady=(0, SPACING['padding']))
        
        # ===== SEÇÃO 1: SELEÇÃO DE PLANILHA =====
        self._criar_secao_planilha()
        
        # ===== SEÇÃO 2: SELEÇÃO DE PASTA DE ARQUIVOS =====
        self._criar_secao_pasta_arquivos()
        
        # ===== SEÇÃO 3: AÇÕES =====
        self._criar_secao_acoes()
        
        # ===== SEÇÃO 4: PROGRESSO =====
        self._criar_secao_progresso()
    
    def _criar_secao_planilha(self):
        """Cria seção de seleção de planilha"""
        frame = ctk.CTkFrame(self, fg_color=COLORS['background'])
        frame.pack(fill="x", pady=SPACING['margin'])
        
        titulo_secao = ctk.CTkLabel(
            frame,
            text="Selecionar Planilha",
            font=FONTS['subtitle'],
            text_color=COLORS['primary']
        )
        titulo_secao.pack(anchor="w", padx=SPACING['padding'], pady=(SPACING['padding'], SPACING['margin']))
        
        # Informação
        info = ctk.CTkLabel(
            frame,
            text="Arquivos suportados: Excel (.xlsx, .xls) ou CSV\nColunas esperadas: CodSAP (buscar) e CodServico (substituir)",
            font=FONTS['small'],
            text_color=COLORS['text_secondary']
        )
        info.pack(anchor="w", padx=SPACING['padding'], pady=(0, SPACING['margin']))
        
        # Frame com botão e label
        container = ctk.CTkFrame(frame, fg_color="transparent")
        container.pack(fill="x", padx=SPACING['padding'], pady=(0, SPACING['padding']))
        
        self.btn_selecionar_planilha = ctk.CTkButton(
            container,
            text="📊 Selecionar Planilha",
            command=self._selecionar_planilha,
            width=250,
            height=40,
            font=FONTS['button'],
            fg_color=COLORS['primary'],
            hover_color=COLORS['hover']
        )
        self.btn_selecionar_planilha.pack(side="left", padx=(0, SPACING['margin']))
        
        # Label com nome do arquivo
        self.label_planilha = ctk.CTkLabel(
            container,
            text="Nenhuma planilha selecionada",
            font=FONTS['input'],
            text_color=COLORS['text_secondary']
        )
        self.label_planilha.pack(side="left", fill="x", expand=True)
    
    def _criar_secao_pasta_arquivos(self):
        """Cria seção de seleção de pasta com arquivos"""
        frame = ctk.CTkFrame(self, fg_color=COLORS['background'])
        frame.pack(fill="x", pady=SPACING['margin'])
        
        titulo_secao = ctk.CTkLabel(
            frame,
            text="Pasta com Arquivos Word",
            font=FONTS['subtitle'],
            text_color=COLORS['primary']
        )
        titulo_secao.pack(anchor="w", padx=SPACING['padding'], pady=(SPACING['padding'], SPACING['margin']))
        
        # Frame com botão e label
        container = ctk.CTkFrame(frame, fg_color="transparent")
        container.pack(fill="x", padx=SPACING['padding'], pady=(0, SPACING['padding']))
        
        self.btn_selecionar_pasta_arquivos = ctk.CTkButton(
            container,
            text="📁 Selecionar Pasta",
            command=self._selecionar_pasta_arquivos,
            width=250,
            height=40,
            font=FONTS['button'],
            fg_color=COLORS['primary'],
            hover_color=COLORS['hover']
        )
        self.btn_selecionar_pasta_arquivos.pack(side="left", padx=(0, SPACING['margin']))
        
        # Label com nome da pasta
        self.label_pasta_arquivos = ctk.CTkLabel(
            container,
            text="Nenhuma pasta selecionada",
            font=FONTS['input'],
            text_color=COLORS['text_secondary']
        )
        self.label_pasta_arquivos.pack(side="left", fill="x", expand=True)
    
    def _criar_secao_acoes(self):
        """Cria seção de botões de ação"""
        frame = ctk.CTkFrame(self, fg_color=COLORS['background'])
        frame.pack(fill="x", pady=SPACING['margin'])
        
        titulo_secao = ctk.CTkLabel(
            frame,
            text="Processar",
            font=FONTS['subtitle'],
            text_color=COLORS['primary']
        )
        titulo_secao.pack(anchor="w", padx=SPACING['padding'], pady=(SPACING['padding'], SPACING['margin']))
        
        container = ctk.CTkFrame(frame, fg_color="transparent")
        container.pack(fill="x", padx=SPACING['padding'], pady=(0, SPACING['padding']))
        
        # Botão processar
        self.btn_processar = ctk.CTkButton(
            container,
            text="🔄 Processar Lote",
            command=self._processar_lote,
            width=250,
            height=45,
            font=FONTS['button'],
            fg_color=COLORS['success'],
            hover_color=('#22863A' if COLORS['success'] == '#28A745' else COLORS['success']),
            state="disabled"
        )
        self.btn_processar.pack(side="left", padx=(0, SPACING['margin']))
        
        # Informação
        self.label_info = ctk.CTkLabel(
            container,
            text="Selecione a planilha e pasta com arquivos",
            font=FONTS['small'],
            text_color=COLORS['text_secondary']
        )
        self.label_info.pack(side="left", fill="x", expand=True)
    
    def _criar_secao_progresso(self):
        """Cria seção de progresso"""
        frame = ctk.CTkFrame(self, fg_color=COLORS['background'])
        frame.pack(fill="x", pady=SPACING['margin'])
        
        titulo_secao = ctk.CTkLabel(
            frame,
            text="Progresso",
            font=FONTS['subtitle'],
            text_color=COLORS['primary']
        )
        titulo_secao.pack(anchor="w", padx=SPACING['padding'], pady=(SPACING['padding'], SPACING['margin']))
        
        container = ctk.CTkFrame(frame, fg_color="transparent")
        container.pack(fill="x", padx=SPACING['padding'], pady=(0, SPACING['padding']))
        
        # Barra de progresso
        self.progress_bar = ctk.CTkProgressBar(
            container,
            height=30,
            fg_color=COLORS['border'],
            progress_color=COLORS['primary']
        )
        self.progress_bar.pack(fill="x", pady=(0, SPACING['margin']))
        self.progress_bar.set(0)
        
        # Label de progresso
        self.label_progresso = ctk.CTkLabel(
            container,
            text="Aguardando processamento...",
            font=FONTS['small'],
            text_color=COLORS['text_secondary']
        )
        self.label_progresso.pack(anchor="w")
    
    def _selecionar_planilha(self):
        """Seleciona a planilha com mapeamento"""
        arquivo = filedialog.askopenfilename(
            title="Selecionar Planilha",
            filetypes=[
                ("Arquivos Excel", "*.xlsx"),
                ("Arquivos Excel 97-2003", "*.xls"),
                ("Arquivos CSV", "*.csv"),
                ("Todos os arquivos", "*.*")
            ],
            parent=self
        )
        
        if arquivo:
            try:
                self.arquivo_planilha = arquivo
                self.dados_planilha = self._carregar_planilha(arquivo)
                
                nome_arquivo = Path(arquivo).name
                self.label_planilha.configure(text=f"✓ {nome_arquivo} ({len(self.dados_planilha)} linhas)")
                self._validar_entrada()
            except Exception as e:
                messagebox.showerror(
                    "Erro ao carregar planilha",
                    f"Erro ao carregar arquivo:\n{str(e)}"
                )
    
    def _carregar_planilha(self, caminho):
        """Carrega a planilha e extrai os dados"""
        try:
            if PANDAS_DISPONIVEL:
                # Usar pandas se disponível
                if caminho.endswith('.csv'):
                    df = pd.read_csv(caminho)
                else:
                    df = pd.read_excel(caminho)
                
                # Encontrar as colunas de busca e substituição (ignorar outras)
                col_busca = None
                col_subst = None
                
                for col in df.columns:
                    col_lower = col.lower()
                    if 'codsap' in col_lower:
                        col_busca = col
                    elif 'codservico' in col_lower or 'cod_servico' in col_lower:
                        col_subst = col
                
                if not col_busca or not col_subst:
                    raise ValueError("Planilha deve conter colunas 'CodSAP' e 'CodServico'")
                
                # Selecionar apenas as duas colunas necessárias
                df_filtrado = df[[col_subst, col_busca]]
                
                dados = []
                for idx, row in df_filtrado.iterrows():
                    busca = str(row[col_busca]).strip()
                    subst = str(row[col_subst]).strip()
                    if busca and subst and busca != 'nan' and subst != 'nan':
                        dados.append((busca, subst))
                
                return dados
            else:
                raise ImportError("pandas não disponível")
        
        except Exception as e:
            raise Exception(f"Erro ao carregar planilha: {str(e)}")
    
    def _selecionar_pasta_arquivos(self):
        """Seleciona a pasta com os arquivos Word"""
        pasta = filedialog.askdirectory(
            title="Selecionar Pasta com Arquivos Word",
            parent=self
        )
        
        if pasta:
            self.pasta_arquivos = pasta
            # Contar arquivos .docx
            arquivos_docx = list(Path(pasta).glob("*.docx"))
            nome_pasta = Path(pasta).name
            self.label_pasta_arquivos.configure(
                text=f"✓ {nome_pasta} ({len(arquivos_docx)} arquivos)"
            )
            self._validar_entrada()
    
    def _validar_entrada(self):
        """Valida se todos os campos estão preenchidos"""
        planilha_ok = self.arquivo_planilha is not None and len(self.dados_planilha) > 0
        pasta_arquivos_ok = self.pasta_arquivos is not None
        
        if planilha_ok and pasta_arquivos_ok:
            self.btn_processar.configure(state="normal")
        else:
            self.btn_processar.configure(state="disabled")
    
    def _processar_lote(self):
        """Inicia o processamento em lote"""
        if not DOCX_DISPONIVEL:
            messagebox.showerror(
                "Erro",
                "Biblioteca python-docx não está instalada"
            )
            return
        
        # Iniciar processamento em thread
        self.processamento_ativo = True
        self.btn_processar.configure(state="disabled", text="⏳ Processando...")
        
        thread = threading.Thread(
            target=self._executar_processamento_lote,
            daemon=True
        )
        thread.start()
    
    def _executar_processamento_lote(self):
        """Executa o processamento de todos os arquivos"""
        try:
            # Obter lista de arquivos
            arquivos = sorted(Path(self.pasta_arquivos).glob("*.docx"))
            total_arquivos = len(arquivos)
            
            if total_arquivos == 0:
                self.after(0, lambda: messagebox.showwarning(
                    "Aviso",
                    "Nenhum arquivo .docx encontrado na pasta selecionada"
                ))
                self.after(0, self._procesamento_erro)
                return
            
            arquivos_processados = 0
            arquivos_desconsiderados = []
            erros = []
            
            # Processar cada arquivo
            for idx, arquivo in enumerate(arquivos):
                try:
                    # Carregar documento
                    doc = Document(str(arquivo))
                    
                    # Rastrear se foi feita alguma substituição neste arquivo
                    arquivo_teve_substituicao = False
                    
                    # OTIMIZAÇÃO: Fazer um único pass percorrendo cada run/célula
                    # e aplicar TODAS as substituições nesse run, em vez de fazer
                    # um pass para cada número (que seria 5000x mais lento)
                    
                    WNS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'

                    def substituir_em_run(run):
                        """Substitui nos <w:t> sem destruir imagens ou outros elementos do run"""
                        alterou = False
                        for elem in run._r.iter(WNS):
                            if elem.text:
                                for numero_busca, numero_substituir in self.dados_planilha:
                                    if numero_busca in elem.text:
                                        elem.text = elem.text.replace(numero_busca, numero_substituir)
                                        alterou = True
                        return alterou

                    # Em parágrafos
                    for paragrafo in doc.paragraphs:
                        for run in paragrafo.runs:
                            if substituir_em_run(run):
                                arquivo_teve_substituicao = True
                    
                    # Em tabelas
                    for tabela in doc.tables:
                        for linha in tabela.rows:
                            for celula in linha.cells:
                                for paragrafo in celula.paragraphs:
                                    for run in paragrafo.runs:
                                        if substituir_em_run(run):
                                            arquivo_teve_substituicao = True
                    
                    # Salvar arquivo apenas se houve substituição
                    if arquivo_teve_substituicao:
                        doc.save(str(arquivo))
                        arquivos_processados += 1
                    else:
                        # Arquivo desconsiderado (nenhum número encontrado)
                        arquivos_desconsiderados.append(arquivo.name)
                
                except Exception as e:
                    erros.append(f"{arquivo.name}: {str(e)}")
                
                # Atualizar progresso
                progresso = (idx + 1) / total_arquivos
                self.after(0, self._atualizar_progresso, progresso, idx + 1, total_arquivos, len(erros))
            
            # Processamento concluído
            self.after(0, self._procesamento_concluido, arquivos_processados, total_arquivos, arquivos_desconsiderados, erros)
        
        except Exception as e:
            self.after(0, lambda: messagebox.showerror(
                "Erro no Processamento",
                f"Erro ao processar lote:\n{str(e)}\n\n{traceback.format_exc()}"
            ))
            self.after(0, self._procesamento_erro)
    
    def _atualizar_progresso(self, progresso, atual, total, erros):
        """Atualiza a barra de progresso"""
        self.progress_bar.set(progresso)
        msg = f"Processando: {atual}/{total} arquivos"
        if erros > 0:
            msg += f" ({erros} erro(s))"
        self.label_progresso.configure(text=msg)
    
    def _procesamento_concluido(self, arquivos_processados, total, arquivos_desconsiderados, erros):
        """Atualiza interface após sucesso"""
        self.processamento_ativo = False
        self.btn_processar.configure(state="normal", text="🔄 Processar Lote")
        
        # Montar mensagem de resultado
        mensagem = f"Processamento concluído!\n\n" \
                   f"{'='*25}\n"
        
        # Arquivos processados
        mensagem += f"✓ ARQUIVOS PROCESSADOS: {arquivos_processados}\n"
        
        # Arquivos desconsiderados
        desconsiderados_count = len(arquivos_desconsiderados)
        mensagem += f"⊘ ARQUIVOS DESCONSIDERADOS: {desconsiderados_count}\n"
        
        if desconsiderados_count > 0:
            mensagem += f"\nArquivos sem nenhum número da planilha:\n"
            for idx, arquivo in enumerate(arquivos_desconsiderados, 1):
                mensagem += f"{idx}. {arquivo}\n"
                # Limitar a 20 arquivos na exibição
                if idx >= 20:
                    if desconsiderados_count > 20:
                        mensagem += f"\n... e mais {desconsiderados_count - 20} arquivo(s)"
                    break
        
        mensagem += f"\n{'='*25}\n"
        mensagem += f"Total de arquivos analisados: {total}\n"
        
        # Mostrar erros se houver
        if erros:
            mensagem += f"\n\n❌ ERROS ENCONTRADOS ({len(erros)}):\n"
            for erro in erros[:10]:  # Mostrar apenas os 10 primeiros erros
                mensagem += f"• {erro}\n"
            if len(erros) > 10:
                mensagem += f"\n... e mais {len(erros) - 10} erro(s)"
        
        messagebox.showinfo("Sucesso", mensagem)
        
        # Gerar arquivo de relatório
        if self.pasta_arquivos:
            self._gerar_relatorio(arquivos_processados, total, arquivos_desconsiderados, erros)
        
        # Atualizar label de progresso
        status_text = f"✓ {arquivos_processados} arquivos processados"
        if desconsiderados_count > 0:
            status_text += f" | ⊘ {desconsiderados_count} desconsiderados"
        
        self.label_progresso.configure(
            text=status_text,
            text_color=COLORS['success']
        )
    
    def _gerar_relatorio(self, arquivos_processados, total, arquivos_desconsiderados, erros):
        """Gera arquivo de relatório em .txt"""
        try:
            # Montar conteúdo do relatório
            relatorio = f"RELATÓRIO DE PROCESSAMENTO - ALTERAR NÚMEROS EM LOTE\n"
            relatorio += f"{'='*60}\n\n"
            
            # Data e hora
            data_hora = datetime.now().strftime("%d/%m/%Y às %H:%M:%S")
            relatorio += f"Data/Hora: {data_hora}\n\n"
            
            # Resumo
            relatorio += f"{'='*60}\n"
            relatorio += f"RESUMO\n"
            relatorio += f"{'='*60}\n"
            relatorio += f"✓ ARQUIVOS PROCESSADOS: {arquivos_processados}\n"
            relatorio += f"⊘ ARQUIVOS DESCONSIDERADOS: {len(arquivos_desconsiderados)}\n"
            relatorio += f"Total de arquivos analisados: {total}\n\n"
            
            # Detalhes dos desconsiderados
            if len(arquivos_desconsiderados) > 0:
                relatorio += f"{'='*60}\n"
                relatorio += f"ARQUIVOS DESCONSIDERADOS\n"
                relatorio += f"{'='*60}\n"
                relatorio += f"Arquivos sem nenhum número da planilha:\n\n"
                for idx, arquivo in enumerate(arquivos_desconsiderados, 1):
                    relatorio += f"{idx}. {arquivo}\n"
                relatorio += f"\n"
            
            # Erros
            if erros:
                relatorio += f"{'='*60}\n"
                relatorio += f"ERROS ENCONTRADOS ({len(erros)})\n"
                relatorio += f"{'='*60}\n"
                for erro in erros:
                    relatorio += f"• {erro}\n"
                relatorio += f"\n"
            
            # Informações adicionais
            relatorio += f"{'='*60}\n"
            relatorio += f"INFORMAÇÕES\n"
            relatorio += f"{'='*60}\n"
            relatorio += f"Pasta com arquivos: {self.pasta_arquivos}\n"
            relatorio += f"Planilha utilizada: {Path(self.arquivo_planilha).name}\n"
            relatorio += f"Total de mapeamentos (CodSAP - CodServico): {len(self.dados_planilha)}\n"
            
            # Salvar arquivo na pasta de origem
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            nome_relatorio = f"relatorio_processamento_{timestamp}.txt"
            caminho_relatorio = Path(self.pasta_arquivos) / nome_relatorio
            
            with open(str(caminho_relatorio), 'w', encoding='utf-8') as f:
                f.write(relatorio)
            
            # Mensagem de confirmação
            messagebox.showinfo(
                "Relatório Gerado",
                f"Arquivo de relatório criado com sucesso:\n\n{nome_relatorio}"
            )
        
        except Exception as e:
            messagebox.showerror(
                "Erro ao Gerar Relatório",
                f"Erro ao criar arquivo de relatório:\n{str(e)}"
            )
    
    def _procesamento_erro(self):
        """Atualiza interface após erro"""
        self.processamento_ativo = False
        self.btn_processar.configure(state="normal", text="🔄 Processar Lote")
        self.label_progresso.configure(
            text="✗ Erro no processamento",
            text_color=COLORS['error']
        )
